import streamlit as st
import PyPDF2
import docx
import pytesseract
from PIL import Image
import hashlib
from datetime import datetime
import os
import cv2
import numpy as np
import fitz  # PyMuPDF
import io
from sentence_transformers import SentenceTransformer


class DocumentProcessor:
    def __init__(self):
        self.supported_types = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'txt': self._process_txt,
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image
        }

        # Advanced processing options
        self.ocr_engines = {
            'tesseract': self._tesseract_ocr,
            'paddleocr': self._paddle_ocr  # If available
        }

        # Initialize sentence transformer for advanced chunking
        try:
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        except:
            self.sentence_model = None
            st.warning("Sentence transformer not available. Using basic chunking.")

    def process_file(self, uploaded_file, use_advanced_processing=False, ocr_enabled=True, table_extraction=True):
        """Process uploaded file with optional advanced features"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()

            if file_extension not in self.supported_types:
                st.error(f"❌ Unsupported file type: {file_extension}")
                return None

            # Generate unique document ID
            file_content = uploaded_file.read()
            doc_id = hashlib.md5(file_content).hexdigest()
            uploaded_file.seek(0)  # Reset file pointer

            # Extract text based on file type and processing mode
            if use_advanced_processing and file_extension == 'pdf':
                text_content = self._process_pdf_advanced(uploaded_file, ocr_enabled, table_extraction)
            else:
                text_content = self.supported_types[file_extension](uploaded_file)

            if not text_content or not text_content.strip():
                st.warning(f"⚠️ No text content extracted from {uploaded_file.name}")
                return None

            # Choose chunking method
            if use_advanced_processing and self.sentence_model:
                chunks = self._advanced_chunk_text(text_content)
            else:
                chunks = self._chunk_text(text_content)

            # Create document data structure
            doc_data = {
                'id': doc_id,
                'name': uploaded_file.name,
                'type': file_extension,
                'content': text_content,
                'size': len(file_content),
                'upload_date': datetime.now(),
                'chunks': chunks,
                'metadata': {
                    'word_count': len(text_content.split()),
                    'char_count': len(text_content),
                    'pages': self._estimate_pages(text_content),
                    'processing_mode': 'advanced' if use_advanced_processing else 'basic',
                    'ocr_enabled': ocr_enabled,
                    'table_extraction': table_extraction
                }
            }

            return doc_data

        except Exception as e:
            st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
            return None

    def _process_pdf(self, uploaded_file):
        """Extract text from PDF file using PyPDF2"""
        try:
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = ""

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

            return text_content
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return ""

    def _process_pdf_advanced(self, uploaded_file, ocr_enabled=True, table_extraction=True):
        """Advanced PDF processing with OCR for images and tables"""
        try:
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            full_text = ""

            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]

                # Extract regular text
                page_text = page.get_text()

                # Extract images and apply OCR if enabled
                if ocr_enabled:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            pix = fitz.Pixmap(pdf_document, xref)

                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_data = pix.tobytes("png")
                                img_text = self._ocr_image(img_data)
                                if img_text.strip():
                                    page_text += f"\n[Image {img_index + 1} OCR]: {img_text}\n"

                            pix = None
                        except Exception as e:
                            st.warning(f"Error processing image {img_index + 1} on page {page_num + 1}: {str(e)}")

                # Extract tables if enabled
                if table_extraction:
                    try:
                        tables = self._extract_tables_from_page(page)
                        for table_index, table in enumerate(tables):
                            page_text += f"\n[Table {table_index + 1}]: {table}\n"
                    except Exception as e:
                        st.warning(f"Error extracting tables from page {page_num + 1}: {str(e)}")

                full_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"

            pdf_document.close()
            return full_text

        except Exception as e:
            st.error(f"Error in advanced PDF processing: {str(e)}")
            return ""

    def _extract_tables_from_page(self, page):
        """Extract tables from PDF page"""
        tables = []
        try:
            # Get table data using PyMuPDF
            table_list = page.find_tables()

            for table in table_list:
                table_data = table.extract()
                if table_data:
                    # Convert table to text format
                    table_text = ""
                    for row in table_data:
                        if row:  # Skip empty rows
                            table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                    tables.append(table_text)

        except Exception as e:
            # Fallback: try to extract table-like structures from text
            page_text = page.get_text()
            tables = self._extract_table_like_structures(page_text)

        return tables

    def _extract_table_like_structures(self, text):
        """Extract table-like structures from text"""
        tables = []
        lines = text.split('\n')

        # Look for lines with multiple columns (separated by spaces or tabs)
        potential_table_lines = []
        for line in lines:
            # Check if line has multiple columns (3+ words with significant spacing)
            words = line.split()
            if len(words) >= 3:
                # Check for consistent spacing patterns
                if '\t' in line or '  ' in line:  # Tab or multiple spaces
                    potential_table_lines.append(line)

        if len(potential_table_lines) >= 2:  # At least 2 rows for a table
            table_text = '\n'.join(potential_table_lines)
            tables.append(table_text)

        return tables

    def _ocr_image(self, img_data):
        """Apply OCR to image data"""
        try:
            # Convert to PIL Image
            image = Image.open(io.BytesIO(img_data))

            # Preprocess image for better OCR
            image = self._preprocess_image(image)

            # Apply OCR with optimized config
            text = pytesseract.image_to_string(image,
                                               config='--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,!?-')
            return text.strip()
        except Exception as e:
            return f"OCR Error: {str(e)}"

    def _preprocess_image(self, image):
        """Preprocess image for better OCR results"""
        try:
            # Convert PIL to OpenCV
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

            # Convert to grayscale
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)

            # Apply denoising
            denoised = cv2.fastNlMeansDenoising(gray)

            # Apply thresholding
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

            # Convert back to PIL
            return Image.fromarray(thresh)
        except Exception as e:
            st.warning(f"Error in image preprocessing: {str(e)}")
            return image

    def _tesseract_ocr(self, image):
        """Tesseract OCR implementation"""
        try:
            return pytesseract.image_to_string(image)
        except Exception as e:
            return f"Tesseract OCR Error: {str(e)}"

    def _paddle_ocr(self, image):
        """PaddleOCR implementation (if available)"""
        try:
            # This would require PaddleOCR installation
            # from paddleocr import PaddleOCR
            # ocr = PaddleOCR(use_angle_cls=True, lang='en')
            # result = ocr.ocr(np.array(image))
            # return ' '.join([line[1][0] for line in result[0]])
            return "PaddleOCR not available"
        except Exception as e:
            return f"PaddleOCR Error: {str(e)}"

    def _process_docx(self, uploaded_file):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(uploaded_file)
            text_content = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content += " | ".join(row_text) + "\n"

            return text_content
        except Exception as e:
            st.error(f"Error processing DOCX: {str(e)}")
            return ""

    def _process_txt(self, uploaded_file):
        """Extract text from TXT file"""
        try:
            text_content = uploaded_file.read().decode('utf-8')
            return text_content
        except UnicodeDecodeError:
            try:
                uploaded_file.seek(0)
                text_content = uploaded_file.read().decode('latin-1')
                return text_content
            except Exception as e:
                st.error(f"Error processing TXT file: {str(e)}")
                return ""

    def _process_image(self, uploaded_file):
        """Extract text from image using OCR"""
        try:
            image = Image.open(uploaded_file)

            # Preprocess image for better OCR
            processed_image = self._preprocess_image(image)

            # Apply OCR
            text_content = pytesseract.image_to_string(processed_image)

            if not text_content.strip():
                st.warning("No text detected in image")
                return ""

            return text_content
        except Exception as e:
            st.error(f"Error processing image with OCR: {str(e)}")
            return ""

    def _chunk_text(self, text, chunk_size=1000, overlap=200):
        """Split text into overlapping chunks (basic method)"""
        words = text.split()
        chunks = []

        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append({
                    'content': chunk,
                    'start_index': i,
                    'end_index': min(i + chunk_size, len(words)),
                    'chunk_id': len(chunks)
                })

        return chunks

    def _advanced_chunk_text(self, text, chunk_size=1000, overlap=200):
        """Advanced chunking with sentence boundary preservation"""
        try:
            # Sentence-aware chunking
            sentences = text.split('.')
            chunks = []
            current_chunk = ""

            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue

                # Check if adding this sentence would exceed chunk size
                if len(current_chunk + sentence + ".") < chunk_size:
                    current_chunk += sentence + "."
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + "."

            # Add the last chunk
            if current_chunk:
                chunks.append(current_chunk.strip())

            # Add overlapping context
            overlapped_chunks = []
            for i, chunk in enumerate(chunks):
                enhanced_chunk = chunk

                # Add previous context
                if i > 0:
                    prev_context = chunks[i - 1][-overlap:] if len(chunks[i - 1]) > overlap else chunks[i - 1]
                    enhanced_chunk = prev_context + " " + enhanced_chunk

                # Add next context
                if i < len(chunks) - 1:
                    next_context = chunks[i + 1][:overlap] if len(chunks[i + 1]) > overlap else chunks[i + 1]
                    enhanced_chunk = enhanced_chunk + " " + next_context

                # Find position in original text
                start_pos = text.find(chunk)
                end_pos = start_pos + len(chunk) if start_pos != -1 else len(text)

                overlapped_chunks.append({
                    'content': enhanced_chunk,
                    'original_chunk': chunk,
                    'chunk_id': i,
                    'start_pos': max(0, start_pos),
                    'end_pos': min(len(text), end_pos)
                })

            return overlapped_chunks

        except Exception as e:
            st.warning(f"Error in advanced chunking, falling back to basic: {str(e)}")
            return self._chunk_text(text, chunk_size, overlap)

    def _estimate_pages(self, text):
        """Estimate number of pages based on text length"""
        words_per_page = 250
        word_count = len(text.split())
        return max(1, word_count // words_per_page)

    def get_processing_stats(self, doc_data):
        """Get detailed processing statistics"""
        stats = {
            'document_name': doc_data['name'],
            'file_size': f"{doc_data['size'] / 1024:.1f} KB",
            'word_count': doc_data['metadata']['word_count'],
            'character_count': doc_data['metadata']['char_count'],
            'estimated_pages': doc_data['metadata']['pages'],
            'chunks_created': len(doc_data['chunks']),
            'processing_mode': doc_data['metadata'].get('processing_mode', 'basic'),
            'ocr_enabled': doc_data['metadata'].get('ocr_enabled', False),
            'table_extraction': doc_data['metadata'].get('table_extraction', False)
        }
        return stats


# Advanced Document Processor Class (Alternative implementation)
class AdvancedDocumentProcessor(DocumentProcessor):
    """Extended version with additional advanced features"""

    def __init__(self):
        super().__init__()
        self.advanced_features_enabled = True

    def process_file_advanced(self, uploaded_file, ocr_enabled=True, table_extraction=True, chunking_strategy="smart"):
        """Advanced processing with all features enabled"""
        return self.process_file(
            uploaded_file,
            use_advanced_processing=True,
            ocr_enabled=ocr_enabled,
            table_extraction=table_extraction
        )

    def batch_process_files(self, uploaded_files, progress_callback=None):
        """Process multiple files with progress tracking"""
        processed_docs = []
        total_files = len(uploaded_files)

        for i, uploaded_file in enumerate(uploaded_files):
            if progress_callback:
                progress_callback(i, total_files, f"Processing {uploaded_file.name}")

            doc_data = self.process_file_advanced(uploaded_file)
            if doc_data:
                processed_docs.append(doc_data)

        return processed_docs
